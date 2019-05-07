import openpyxl
from openpyxl.styles import Color, PatternFill, Font, Border
from openpyxl.styles import colors
from openpyxl.cell import Cell
import win32com.client
import os
import docx
from docx import Document
import datetime
import  smtplib
from win32com.client import Dispatch
from openpyxl.styles.borders import Border, Side
excel = win32com.client.Dispatch("Excel.Application")
wb = openpyxl.load_workbook('students.xlsx')



sheet_1 = wb.active

border = Border(left=Side(style='medium'), right=Side(style='medium'), top=Side(style='medium'), bottom=Side(style='medium'))
groups_list = []
groups_id = []

for i in range(2, sheet_1.max_row):
    if str(sheet_1.cell(row=i, column=3).value) not in groups_list:
        groups_list.append(str(sheet_1.cell(row=i, column=3).value))
        groups_id.append(i)

ws1 = wb.create_sheet("5.01")
ws2 = wb.create_sheet("5.02")
ws3 = wb.create_sheet("5.03")
for i in range(1, 7):
    ws1.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value
    ws1.cell(row=1, column=i).border = border
for i in range(1, 7):
    ws2.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value
    ws2.cell(row=1, column=i).border = border
for i in range(1, 7):
    ws3.cell(row=1, column=i).value = sheet_1.cell(row=1, column=i).value
    ws3.cell(row=1, column=i).border = border

wb.save("students_edit.xlsx")
for i in range(2, groups_id[1] - 1):
    ws1.cell(row=i, column=1).value = i - 1
    ws1.cell(row=i, column=1).border = border
for i in range(1, int(groups_id[1]) - 1):
    for j in range(2, 7):
        ws1.cell(row=i, column=j).value = sheet_1.cell(row=i, column=j).value
        ws1.cell(row=i, column=j).border = border
wb.save("students_edit.xlsx")

for i in range(2, groups_id[1] - 1):
    ws2.cell(row=i, column=1).value = i - 1
    ws2.cell(row=i, column=1).border = border

for i in range(int(groups_id[1]), int(groups_id[2])):
    for j in range(2, 7):
        ws2.cell(row=i - 6, column=j).value = sheet_1.cell(row=i, column=j).value
        ws2.cell(row=i-6, column=j).border = border
wb.save("students_edit.xlsx")

for i in range(2, groups_id[1] - 1):
    ws3.cell(row=i, column=1).value = i - 1
    ws3.cell(row=i, column=1).border = border

for i in range(int(groups_id[2]), sheet_1.max_row):
    for j in range(2, 7):
        ws3.cell(row=i - 11, column=j).value = sheet_1.cell(row=i, column=j).value
        ws3.cell(row=i - 11, column=j).border = border


ws1.cell(row=1, column=7).value = "Mark"
ws1.cell(row=1, column=7).border = border
ws2.cell(row=1, column=7).value = "Mark"
ws2.cell(row=1, column=7).border = border
ws3.cell(row=1, column=7).value = "Mark"
ws3.cell(row=1, column=7).border = border

value = 0
for i in range(2, ws1.max_row + 1):

    value = 0
    for j in range(4, 7):
        value = value + ws1.cell(row=i, column=j).value

        ws1.cell(row=i, column=7).value = round(value / 3)
        ws1.cell(row=i, column=7).border = border

for i in range(2, ws2.max_row + 1):
    value = 0
    for j in range(4, 7):
        value = value + ws2.cell(row=i, column=j).value
        ws2.cell(row=i, column=7).value = round(value / 3)
        ws2.cell(row=i, column=7).border = border

for i in range(2, ws3.max_row + 1):
    value = 0
    for j in range(4, 7):
        value = value + ws3.cell(row=i, column=j).value
        ws3.cell(row=i, column=7).value = round(value / 3)
        ws3.cell(row=i, column=7).border = border


wb.save("students_edit.xlsx")

wb.close()
temp = 0

wb = excel.Workbooks.Open(os.path.join(os.getcwd(), "students_edit.xlsx"))
ws1 = wb.Worksheets("5.01")
ws1.Range('A2:G6').Sort(Key1=ws1.Range('G6'), Order1=2, Orientation=1)
ws2 = wb.Worksheets("5.02")
ws2.Range('A2:G6').Sort(Key1=ws2.Range('G6'), Order1=2, Orientation=1)
ws3 = wb.Worksheets("5.03")
ws3.Range('A2:G6').Sort(Key1=ws3.Range('G6'), Order1=2, Orientation=1)

wb.Save()
excel.Application.Quit()

redFill = PatternFill(start_color='FFFF0000', end_color='FFFF0000', fill_type='solid')

wb = openpyxl.load_workbook('students_edit.xlsx')
ws1 = wb["5.01"]
ws2 = wb["5.02"]
ws3 = wb["5.03"]
badstudents = []
for i in range(2, ws1.max_row + 1):
    if ws1.cell(row=i, column=7).value <= 60:
        ws1.cell(row=i, column=7).fill = redFill
        badstudents.append(ws1.cell(row=i, column=2).value)
for i in range(2, ws2.max_row + 1):
    if ws2.cell(row=i, column=7).value <= 60:
        ws2.cell(row=i, column=7).fill = redFill
        badstudents.append(ws2.cell(row=i, column=2).value)
for i in range(2, ws3.max_row + 1):
    if ws3.cell(row=i, column=7).value <= 60:
        ws3.cell(row=i, column=7).fill = redFill
        badstudents.append(ws3.cell(row=i, column=2).value)

wb.save("students_edit.xlsx")

ws4 = wb.create_sheet("Боржники")
cntr = 0
ws4.cell(row=1, column=1).value = "Боржники"
ws4.cell(row=1, column=1).border = border
for i in range(2, len(badstudents) + 2):
    ws4.cell(row=i, column=1).value = badstudents[cntr]
    ws4.cell(row=i, column=1).border = border
    cntr = cntr+1

wb.save("students_edit.xlsx")
document = Document()
cntr = 0

paragraph = document.add_paragraph("Боржники РАПОРТ")
for i in range(0,len(badstudents)):
    paragraph = document.add_paragraph(badstudents[i])
document.save("borjniki.docx")
